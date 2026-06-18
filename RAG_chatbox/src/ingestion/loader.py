from __future__ import annotations

import csv
import json
from pathlib import Path

from ..utils.schemas import Document


TEXT_EXTENSIONS = {".txt", ".md", ".rst", ".py", ".yaml", ".yml", ".csv", ".json", ".jsonl"}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | {".pdf", ".docx"}


def load_documents(paths: list[str | Path]) -> list[Document]:
    documents: list[Document] = []
    for path in paths:
        resolved = Path(path).expanduser().resolve()
        if resolved.is_dir():
            files = sorted(
                file for file in resolved.rglob("*") if file.is_file() and file.suffix.lower() in SUPPORTED_EXTENSIONS
            )
            for file in files:
                documents.extend(_load_file(file))
        elif resolved.is_file():
            documents.extend(_load_file(resolved))
        else:
            raise FileNotFoundError(f"Path does not exist: {resolved}")
    return [doc for doc in documents if doc.text.strip()]


def _load_file(path: Path) -> list[Document]:
    suffix = path.suffix.lower()
    if suffix == ".csv":
        return [_load_csv(path)]
    if suffix == ".json":
        return _load_json(path)
    if suffix == ".jsonl":
        return _load_jsonl(path)
    if suffix == ".pdf":
        return [_load_pdf(path)]
    if suffix == ".docx":
        return [_load_docx(path)]
    if suffix in TEXT_EXTENSIONS:
        return [Document(text=path.read_text(encoding="utf-8"), source=str(path))]
    raise ValueError(f"Unsupported file type: {path}")


def _load_csv(path: Path) -> Document:
    rows: list[str] = []
    with path.open("r", encoding="utf-8", newline="") as handle:
        reader = csv.DictReader(handle)
        for row in reader:
            rows.append(" | ".join(f"{key}: {value}" for key, value in row.items() if value))
    return Document(text="\n".join(rows), source=str(path), metadata={"format": "csv"})


def _load_json(path: Path) -> list[Document]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if isinstance(data, list):
        return [_document_from_json_item(item, path, idx) for idx, item in enumerate(data)]
    return [_document_from_json_item(data, path, 0)]


def _load_jsonl(path: Path) -> list[Document]:
    documents: list[Document] = []
    with path.open("r", encoding="utf-8") as handle:
        for idx, line in enumerate(handle):
            line = line.strip()
            if not line:
                continue
            documents.append(_document_from_json_item(json.loads(line), path, idx))
    return documents


def _document_from_json_item(item: object, path: Path, idx: int) -> Document:
    if isinstance(item, dict):
        text = item.get("context") or item.get("text") or item.get("content") or json.dumps(item, ensure_ascii=False)
        metadata = {key: value for key, value in item.items() if key not in {"context", "text", "content"}}
    else:
        text = str(item)
        metadata = {}
    return Document(text=str(text), source=f"{path}#{idx}", metadata=metadata)


def _load_pdf(path: Path) -> Document:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ImportError("PDF loading requires pypdf. Install with: pip install pypdf") from exc

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return Document(text="\n\n".join(pages), source=str(path), metadata={"format": "pdf"})


def _load_docx(path: Path) -> Document:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise ImportError("DOCX loading requires python-docx. Install with: pip install python-docx") from exc

    document = DocxDocument(str(path))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    return Document(text=text, source=str(path), metadata={"format": "docx"})
